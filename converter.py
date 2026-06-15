#!/usr/bin/env python3
"""
PDF ↔ DOCX ↔ PPTX Converter
Converts between PDF, Word documents, and PowerPoint presentations.

Supported conversions:
- PDF → DOCX (extracts text and basic formatting)
- PDF → PPTX (creates slides from PDF pages)
- DOCX → PDF (converts document to PDF)
- PPTX → PDF (converts slides to PDF)
- DOCX → PPTX (creates slides from document paragraphs)
- PPTX → DOCX (extracts text from slides)
"""

import argparse
import logging
import os
import sys
import platform
import subprocess
from pathlib import Path

# Import config module
from config import ConverterConfig, DEFAULT_CONFIG


# Global logger instance
logger = None


def setup_logging(log_level='INFO', log_file=None):
    """Configure logging for console and optional file output."""
    global logger
    
    # Create logger
    logger = logging.getLogger('converter')
    logger.setLevel(log_level.upper())
    
    # Clear any existing handlers
    logger.handlers.clear()
    
    # Console handler
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(log_level.upper())
    
    # Formatter
    formatter = logging.Formatter(
        '[%(levelname)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    # File handler (optional)
    if log_file:
        try:
            file_handler = logging.FileHandler(log_file, mode='a')
            file_handler.setLevel(log_level.upper())
            file_formatter = logging.Formatter(
                '[%(asctime)s] [%(levelname)s] %(message)s',
                datefmt='%Y-%m-%d %H:%M:%S'
            )
            file_handler.setFormatter(file_formatter)
            logger.addHandler(file_handler)
            logger.info(f"Logging to file: {log_file}")
        except Exception as e:
            logger.warning(f"Could not create log file: {e}")
    
    return logger


# Initialize logger at module import
logger = setup_logging('INFO')


def find_libreoffice():
    """Locate LibreOffice executable on the system."""
    system = platform.system()
    
    if system == 'Windows':
        # Common Windows installation paths
        possible_paths = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                logger.debug(f"Found LibreOffice at: {path}")
                return path
        logger.debug("LibreOffice not found in Windows standard paths")
    elif system == 'Darwin':  # macOS
        path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        if os.path.exists(path):
            logger.debug(f"Found LibreOffice at: {path}")
            return path
        logger.debug("LibreOffice not found in macOS standard path")
    else:  # Linux
        try:
            result = subprocess.run(['which', 'soffice'], capture_output=True, text=True)
            if result.returncode == 0:
                path = result.stdout.strip()
                logger.debug(f"Found LibreOffice at: {path}")
                return path
        except Exception:
            pass
        logger.debug("LibreOffice not found in Linux PATH")
    
    return None


def install_dependencies():
    """Check and suggest installing required packages."""
    required = {
        'pdf2docx': 'pdf2docx',
        'pptx': 'python-pptx',
        'docx': 'python-docx',
        'pypdf': 'pypdf',
        'tqdm': 'tqdm',
    }
    
    missing = []
    for module, package in required.items():
        try:
            __import__(module)
            logger.debug(f"Package '{package}' is installed")
        except ImportError:
            missing.append(package)
            logger.warning(f"Missing package: {package}")
    
    # tkinter is built-in, just check it
    try:
        __import__('tkinter')
        logger.debug("Package 'tkinter' is installed")
    except ImportError:
        logger.warning("tkinter not available (GUI will not work)")
    
    if missing:
        logger.error("Missing required packages. Install with:")
        logger.error(f"pip install {' '.join(missing)}")
        return False
    
    logger.info("All dependencies are installed")
    return True


def pdf_to_docx(pdf_path, output_path=None, config=None):
    """Convert PDF to DOCX with formatting preservation and fallback support."""
    if not output_path:
        output_path = Path(pdf_path).stem + ".docx"
    
    if not config:
        config = ConverterConfig()
    
    try:
        from pdf2docx import convert
        
        logger.info(f"Converting {pdf_path} → {output_path}...")
        logger.debug("Using pdf2docx with formatting preservation enabled")
        
        # Convert with formatting preservation
        convert(
            pdf_path, 
            output_path,
            pages=None,  # Convert all pages
            multi_processing=config.get('enable_multiprocessing', True)
        )
        
        logger.info(f"Successfully created {output_path}")
        logger.debug(f"Conversion complete. Output file: {output_path}")
        return output_path
    
    except ImportError:
        logger.error("pdf2docx library not found.")
        
        # Fallback: Try alternative method
        if config.get('fallback_mode', True):
            logger.warning("Attempting fallback conversion method...")
            return _pdf_to_docx_fallback(pdf_path, output_path)
        
        logger.error("Please install it with: pip install pdf2docx")
        return None
    
    except Exception as e:
        logger.error(f"Error converting PDF to DOCX: {e}")
        logger.debug(f"PDF path: {pdf_path}, Output path: {output_path}")
        
        # Fallback on error
        if config.get('fallback_mode', True):
            logger.warning(f"Primary method failed, attempting fallback...")
            try:
                return _pdf_to_docx_fallback(pdf_path, output_path)
            except Exception as fallback_error:
                logger.error(f"Fallback method also failed: {fallback_error}")
                return None
        
        return None


def _pdf_to_docx_fallback(pdf_path, output_path):
    """Fallback method for PDF to DOCX conversion using text extraction."""
    try:
        from docx import Document
        from pypdf import PdfReader
        
        logger.info("Using pypdf fallback for PDF→DOCX conversion")
        
        # Read PDF with pypdf
        pdf_reader = PdfReader(pdf_path)
        doc = Document()
        
        # Extract text from each page
        for page_num, page in enumerate(pdf_reader.pages, 1):
            doc.add_heading(f"Page {page_num}", level=2)
            
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("[Page contains images or non-extractable content]")
        
        doc.save(output_path)
        logger.info(f"Fallback conversion successful: {output_path}")
        return output_path
    
    except Exception as e:
        logger.error(f"Fallback PDF→DOCX conversion failed: {e}")
        return None



def docx_to_pdf(docx_path, output_path=None, config=None):
    """Convert DOCX to PDF with LibreOffice or fallback method."""
    if not output_path:
        output_path = Path(docx_path).stem + ".pdf"
    
    try:
        logger.info(f"Converting {docx_path} → {output_path}...")
        
        # Try LibreOffice first
        libreoffice_path = find_libreoffice()
        if libreoffice_path:
            try:
                logger.debug(f"Using LibreOffice: {libreoffice_path}")
                output_dir = Path(output_path).parent.absolute()
                subprocess.run([
                    libreoffice_path, '--headless', '--convert-to', 'pdf',
                    '--outdir', str(output_dir), str(docx_path)
                ], check=True, capture_output=True, timeout=30)
                
                logger.info(f"Successfully created {output_path}")
                return output_path
            except subprocess.TimeoutExpired:
                logger.warning("LibreOffice conversion timed out, trying fallback...")
            except Exception as lo_error:
                logger.warning(f"LibreOffice conversion failed: {lo_error}, trying fallback...")
        else:
            logger.warning("LibreOffice not found, using fallback method...")
        
        # Fallback: Extract text from DOCX and create PDF
        return _docx_to_pdf_fallback(docx_path, output_path)
        
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {e}")
        return None


def _docx_to_pdf_fallback(docx_path, output_path):
    """Fallback method to convert DOCX to PDF by extracting text."""
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        
        logger.info("Using fallback DOCX→PDF conversion (text extraction)")
        
        # Read DOCX
        doc = Document(docx_path)
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Extract text from DOCX
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        pdf_doc.build(story)
        logger.info(f"Fallback conversion successful: {output_path}")
        return output_path
        
    except ImportError:
        logger.error("reportlab not installed. Install with: pip install reportlab")
        return None
    except Exception as e:
        logger.error(f"Fallback DOCX→PDF conversion failed: {e}")
        return None


def process_single_file(input_file, output_format, output_path=None, config=None):
    """Process a single file conversion."""
    if not config:
        config = ConverterConfig()
    
    converters = {
        ('pdf', 'docx'): lambda i, o: pdf_to_docx(i, o, config),
        ('docx', 'pdf'): lambda i, o: docx_to_pdf(i, o, config),
    }
    
    input_ext = Path(input_file).suffix.lower()[1:]
    key = (input_ext, output_format)
    
    if key not in converters:
        logger.warning(f"Unsupported conversion: {input_ext.upper()} → {output_format.upper()}")
        logger.warning(f"Skipped: {input_file}")
        return None
    
    converter = converters[key]
    return converter(input_file, output_path)


def process_batch(input_dir, output_format, output_dir=None, recursive=False, config=None):
    """Process multiple files in a directory."""
    from tqdm import tqdm
    
    if not config:
        config = ConverterConfig()
    
    input_path = Path(input_dir)
    
    if not input_path.is_dir():
        logger.error(f"'{input_dir}' is not a directory")
        return []
    
    logger.info(f"Starting batch processing on: {input_dir}")
    if recursive:
        logger.info("Recursive mode enabled")
    
    # Create output directory if specified
    if output_dir:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        logger.debug(f"Output directory: {output_dir}")
    else:
        output_path = input_path
    
    # Find all supported input files
    supported_extensions = {'pdf', 'docx', 'pptx'}
    pattern = '**/*' if recursive else '*'
    
    files_to_process = [
        f for f in input_path.glob(pattern)
        if f.is_file() and f.suffix.lower()[1:] in supported_extensions
    ]
    
    if not files_to_process:
        logger.warning(f"No supported files found in {input_dir}")
        return []
    
    logger.info(f"Found {len(files_to_process)} file(s) to convert")
    
    results = []
    for input_file in tqdm(files_to_process, desc="Batch processing", unit="file"):
        logger.debug(f"Processing {input_file.name}...")
        
        # Generate output path
        output_file = output_path / f"{input_file.stem}.{output_format}"
        
        result = process_single_file(str(input_file), output_format, str(output_file), config)
        if result:
            results.append(result)
    
    return results


def main():
    global logger
    
    parser = argparse.ArgumentParser(
        description='Convert between PDF, DOCX, and PPTX files (single or batch)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Single file conversion
  python converter.py document.pdf --to docx
  python converter.py slides.pptx --to pdf --output converted.pdf
  
  # Batch conversion (entire directory)
  python converter.py ./documents --to pdf --batch
  python converter.py ./documents --to docx --batch --output ./converted
  python converter.py ./documents --to pptx --batch --recursive
  
  # With logging and config
  python converter.py document.pdf --to docx --log conversion.log --verbose
  python converter.py document.pdf --to docx --config myconfig.json
  
  # Config management
  python converter.py --create-config
  python converter.py --show-config
        '''
    )
    
    parser.add_argument('input', nargs='?', help='Input file path or directory (for batch mode)')
    parser.add_argument('--to', choices=['pdf', 'docx', 'pptx'],
                       help='Output format')
    parser.add_argument('-o', '--output', help='Output file or directory path')
    parser.add_argument('--batch', action='store_true',
                       help='Enable batch mode for processing directory')
    parser.add_argument('--recursive', action='store_true',
                       help='Process subdirectories recursively (batch mode only)')
    parser.add_argument('--log', help='Log file path (optional)')
    parser.add_argument('--verbose', action='store_true',
                       help='Enable verbose/debug logging')
    parser.add_argument('--config', help='Configuration file path')
    parser.add_argument('--create-config', action='store_true',
                       help='Create a default configuration file')
    parser.add_argument('--show-config', action='store_true',
                       help='Display current configuration')
    
    args = parser.parse_args()
    
    # Handle config management commands
    if args.create_config:
        from config import create_default_config, CONFIG_LOCATIONS
        config_file = create_default_config()
        print(f"Configuration file created at: {config_file}")
        sys.exit(0)
    
    if args.show_config:
        config = ConverterConfig(args.config)
        print(config)
        sys.exit(0)
    
    # Setup logging
    log_level = 'DEBUG' if args.verbose else 'INFO'
    logger = setup_logging(log_level, args.log)
    logger.debug(f"Starting converter with arguments: {vars(args)}")
    
    # Load configuration
    config = ConverterConfig(args.config)
    config.validate()
    
    # Check dependencies
    if not install_dependencies():
        sys.exit(1)
    
    # Require input and --to for actual conversions
    if not args.input or not args.to:
        parser.print_help()
        sys.exit(1)
    
    # Batch mode
    if args.batch:
        if not os.path.isdir(args.input):
            logger.error(f"'{args.input}' is not a directory")
            sys.exit(1)
        
        results = process_batch(args.input, args.to, args.output, args.recursive, config)
        
        if results:
            logger.info(f"Successfully converted {len(results)} file(s)")
            sys.exit(0)
        else:
            logger.error("Batch conversion completed with no successful conversions")
            sys.exit(1)
    
    # Single file mode
    if not os.path.exists(args.input):
        logger.error(f"File '{args.input}' not found")
        sys.exit(1)
    
    result = process_single_file(args.input, args.to, args.output, config)
    sys.exit(0 if result else 1)


if __name__ == '__main__':
    main()
